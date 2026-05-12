"""Report export: Excel and Word generation."""

import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List


def export_excel(diff_result: Dict = None, enrich_result: Dict = None,
                gsea_result: Dict = None, output_dir: Path = None) -> Dict:
    """Export all results to a multi-sheet Excel workbook."""
    try:
        import openpyxl
        from openpyxl.utils import get_column_letter
    except ImportError:
        return {"error": "openpyxl not installed. Run: pip install openpyxl"}

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    # Differential sheet
    if diff_result:
        ws = wb.create_sheet("Differential_Genes")
        ws.append(["Gene", "log2FC", "P-value", "FDR", "Direction", "Mean_Group1", "Mean_Group2"])
        for g in diff_result.get('all_genes', []):
            ws.append([g.get('gene'), g.get('log2FC'), g.get('pvalue'), g.get('fdr'),
                      g.get('direction'), g.get('mean_group1'), g.get('mean_group2')])

        ws2 = wb.create_sheet("Top_Up_Genes")
        ws2.append(["Gene", "log2FC", "FDR"])
        for g in [g for g in diff_result.get('all_genes', []) if g.get('direction') == 'up'][:100]:
            ws2.append([g.get('gene'), g.get('log2FC'), g.get('fdr')])

        ws3 = wb.create_sheet("Top_Down_Genes")
        ws3.append(["Gene", "log2FC", "FDR"])
        for g in [g for g in diff_result.get('all_genes', []) if g.get('direction') == 'down'][:100]:
            ws3.append([g.get('gene'), g.get('log2FC'), g.get('fdr')])

    # Enrichment sheet
    if enrich_result:
        ws = wb.create_sheet("Enrichment")
        ws.append(["Term", "ID", "Category", "Gene_Count", "P-value", "FDR", "Genes"])
        for t in enrich_result.get('terms', []):
            ws.append([t.get('term'), t.get('id'), t.get('category'),
                      t.get('gene_count'), t.get('pvalue'), t.get('fdr'),
                      ', '.join(t.get('genes', []))])

    # GSEA sheet
    if gsea_result:
        ws = wb.create_sheet("GSEA")
        ws.append(["Term", "ID", "NES", "ES", "P-value", "FDR", "Leading_Edge_Genes"])
        for t in gsea_result.get('terms', []):
            ws.append([t.get('term'), t.get('id'), t.get('nes'), t.get('es'),
                      t.get('pvalue'), t.get('fdr'), ', '.join(t.get('leading_edge', []))])

    # Summary sheet
    ws = wb.create_sheet("Summary")
    ws.append(["Analysis Summary", datetime.now().isoformat()])
    ws.append([])
    if diff_result:
        ws.append(["Differential Analysis"])
        ws.append(["Total genes tested", diff_result.get('n_total')])
        ws.append(["Up-regulated", diff_result.get('n_up')])
        ws.append(["Down-regulated", diff_result.get('n_down')])
        ws.append(["Method", diff_result.get('method')])
        ws.append(["log2FC threshold", diff_result.get('logfc_threshold')])
        ws.append(["FDR threshold", diff_result.get('fdr_threshold')])

    out_dir = Path(output_dir) if output_dir else Path('output')
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"analysis_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    wb.save(str(path))

    return {"files": [{"name": "Excel Results", "path": str(path), "size": f"{path.stat().st_size/1024:.0f} KB"}]}


def export_word(diff_result: Dict = None, enrich_result: Dict = None,
                gsea_result: Dict = None, interpretation: Dict = None,
                output_dir: Path = None) -> Dict:
    """Generate a Word report."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    doc = Document()
    doc.add_heading('Bioinformatics Analysis Report', 0)
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')

    # Data summary
    doc.add_heading('1. Data Overview', level=1)
    if diff_result:
        doc.add_paragraph(f"Total genes tested: {diff_result.get('n_total', 'N/A')}")
        doc.add_paragraph(f"Differentially expressed: {diff_result.get('n_up', 0)} up, {diff_result.get('n_down', 0)} down")
        doc.add_paragraph(f"Method: {diff_result.get('method', 'N/A')}, |log2FC| > {diff_result.get('logfc_threshold', 'N/A')}, FDR < {diff_result.get('fdr_threshold', 'N/A')}")

    # Differential results
    doc.add_heading('2. Differential Expression Analysis', level=1)
    if diff_result:
        doc.add_paragraph(f"Found {diff_result.get('n_sig', 0)} significantly differentially expressed genes.")
        if diff_result.get('top_genes'):
            table = doc.add_table(rows=1, cols=5, style='Light Grid Accent 1')
            for j, h in enumerate(['Gene', 'log2FC', 'P-value', 'FDR', 'Direction']):
                table.rows[0].cells[j].text = h
            for g in diff_result.get('top_genes', [])[:30]:
                row = table.add_row()
                row.cells[0].text = str(g.get('gene', ''))
                row.cells[1].text = f"{g.get('logfc', 0):.3f}"
                row.cells[2].text = f"{g.get('pval', 1):.1e}"
                row.cells[3].text = f"{g.get('fdr', 1):.1e}"
                row.cells[4].text = g.get('direction', '')

    # Enrichment
    doc.add_heading('3. Enrichment Analysis', level=1)
    if enrich_result:
        doc.add_paragraph(f"Found {enrich_result.get('n_terms', 0)} significantly enriched terms.")
        if enrich_result.get('terms'):
            table = doc.add_table(rows=1, cols=4, style='Light Grid Accent 1')
            for j, h in enumerate(['Term', 'Gene Count', 'P-value', 'FDR']):
                table.rows[0].cells[j].text = h
            for t in enrich_result.get('terms', [])[:20]:
                row = table.add_row()
                row.cells[0].text = str(t.get('term', ''))[:80]
                row.cells[1].text = str(t.get('gene_count', ''))
                row.cells[2].text = f"{t.get('pvalue', 1):.1e}"
                row.cells[3].text = f"{t.get('fdr', 1):.1e}"

    # GSEA
    if gsea_result:
        doc.add_heading('4. GSEA Analysis', level=1)
        doc.add_paragraph(f"Found {gsea_result.get('n_terms', 0)} enriched gene sets.")
        if gsea_result.get('terms'):
            table = doc.add_table(rows=1, cols=5, style='Light Grid Accent 1')
            for j, h in enumerate(['Term', 'NES', 'P-value', 'FDR', 'Leading Edge']):
                table.rows[0].cells[j].text = h
            for t in gsea_result.get('terms', [])[:15]:
                row = table.add_row()
                row.cells[0].text = str(t.get('term', ''))[:60]
                row.cells[1].text = f"{t.get('nes', 0):.3f}"
                row.cells[2].text = f"{t.get('pvalue', 1):.1e}"
                row.cells[3].text = f"{t.get('fdr', 1):.1e}"
                row.cells[4].text = ', '.join(t.get('leading_edge', [])[:5])

    # Interpretation
    if interpretation:
        doc.add_heading('5. Result Interpretation', level=1)
        for key in ['summary', 'results', 'discussion']:
            if interpretation.get(key):
                doc.add_heading(key.replace('_', ' ').title(), level=2)
                doc.add_paragraph(interpretation[key])

    # Methods
    doc.add_heading('6. Methods', level=1)
    doc.add_paragraph("Analysis performed using BioInfo Platform. Differential expression analyzed with statistical testing and Benjamini-Hochberg correction. Enrichment analysis using hypergeometric test. GSEA using pre-ranked gene list approach.")

    out_dir = Path(output_dir) if output_dir else Path('output')
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(str(path))

    return {"files": [{"name": "Word Report", "path": str(path), "size": f"{path.stat().st_size/1024:.0f} KB"}]}
